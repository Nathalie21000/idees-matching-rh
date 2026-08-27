import pdfplumber
import docx
import re
import pytesseract


# ============================================================
# EXTRACTION DE TEXTE
# PDF + WORD + OCR PDF SCANNÉ
# ============================================================

def extract_text(file):
    """
    Extrait le texte d'un fichier PDF ou Word (.docx).

    PDF :
    - utilise d'abord pdfplumber ;
    - si aucune texte n'est trouvé, utilise Tesseract OCR.

    DOCX :
    - lit les paragraphes ;
    - lit également les tableaux.
    """

    nom_fichier = getattr(file, "name", "") or ""
    nom_fichier_min = nom_fichier.lower()

    if nom_fichier_min.endswith(".docx"):
        texte = extraire_texte_docx(file)

    elif nom_fichier_min.endswith(".pdf"):
        texte = extraire_texte_pdf(file)

    else:
        texte = ""

    return nettoyer_texte(texte)


# ============================================================
# EXTRACTION PDF
# ============================================================

def extraire_texte_pdf(file):
    """
    Extrait le texte d'un PDF.

    Si pdfplumber ne trouve pas de texte sur une page,
    la page est passée à Tesseract.
    """

    morceaux = []

    try:

        with pdfplumber.open(file) as pdf:

            for page in pdf.pages:

                texte_page = page.extract_text(
                    x_tolerance=2,
                    y_tolerance=3
                )

                if texte_page and texte_page.strip():

                    morceaux.append(texte_page)

                else:

                    texte_ocr = extraire_page_avec_ocr(page)

                    if texte_ocr:
                        morceaux.append(texte_ocr)

    except Exception:
        return ""

    return "\n".join(morceaux)


# ============================================================
# OCR
# ============================================================

def extraire_page_avec_ocr(page):
    """
    Convertit une page PDF en image puis utilise Tesseract
    avec le modèle français.
    """

    try:

        image_page = page.to_image(
            resolution=300
        ).original

        texte = pytesseract.image_to_string(
            image_page,
            lang="fra",
            config="--psm 6"
        )

        return texte or ""

    except Exception:
        return ""


# ============================================================
# EXTRACTION WORD
# ============================================================

def extraire_texte_docx(file):
    """
    Extrait le texte d'un document Word .docx.

    Les paragraphes et les tableaux sont lus.
    """

    morceaux = []

    try:

        document = docx.Document(file)

        # Paragraphes
        for paragraphe in document.paragraphs:

            texte = paragraphe.text.strip()

            if texte:
                morceaux.append(texte)

        # Tableaux
        for table in document.tables:

            for ligne in table.rows:

                cellules = []

                for cellule in ligne.cells:

                    texte_cellule = cellule.text.strip()

                    if texte_cellule:
                        cellules.append(texte_cellule)

                if cellules:

                    morceaux.append(
                        " | ".join(cellules)
                    )

    except Exception:
        return ""

    return "\n".join(morceaux)


# ============================================================
# NETTOYAGE
# ============================================================

def nettoyer_texte(texte):
    """
    Nettoyage léger du texte.
    """

    if not texte:
        return ""

    texte = texte.replace("\r\n", "\n")
    texte = texte.replace("\r", "\n")
    texte = texte.replace("\xa0", " ")

    texte = re.sub(
        r"[ \t]+",
        " ",
        texte
    )

    texte = re.sub(
        r"\n[ \t]*\n[ \t]*\n+",
        "\n\n",
        texte
    )

    return texte.strip()


# ============================================================
# NORMALISATION
# ============================================================

def _normaliser_ligne(ligne):
    """
    Normalise une ligne pour faciliter la reconnaissance
    des libellés malgré les petites erreurs OCR.
    """

    if not ligne:
        return ""

    ligne = ligne.strip().lower()

    ligne = ligne.replace("’", "'")

    ligne = re.sub(
        r"[|¦]+",
        " ",
        ligne
    )

    ligne = re.sub(
        r"\s+",
        " ",
        ligne
    )

    return ligne.strip()


# ============================================================
# LIBELLÉ ENTREPRISE
# ============================================================

def _est_libelle_entreprise(ligne):

    texte = _normaliser_ligne(ligne)

    return bool(
        re.search(
            r"nom\s+de\s+l['’]?entreprise",
            texte,
            re.IGNORECASE
        )
    )


# ============================================================
# LIBELLÉ TÂCHES
# ============================================================

def _est_libelle_taches(ligne):

    texte = _normaliser_ligne(ligne)

    return bool(
        re.search(
            r"liste\s+des\s+t[âa]ches\s+propos[ée]es",
            texte,
            re.IGNORECASE
        )
    )


# ============================================================
# LIBELLÉ POSTE
# ============================================================

def _est_libelle_poste(ligne):

    texte = _normaliser_ligne(ligne)

    # Reconnaissance normale
    if re.search(
        r"intitul[ée]\s+du\s+poste",
        texte,
        re.IGNORECASE
    ):
        return True

    # --------------------------------------------------------
    # Variantes OCR observées sur notre document
    # --------------------------------------------------------

    variantes = [
        "intitule du poste",
        "intitul du poste",
        "intitule poste",
        "intitul poste",
        "linie du poste",
        "linte du poste",
        "linie poste",
        "intitule du poste",
    ]

    for variante in variantes:

        if variante in texte:
            return True

    # --------------------------------------------------------
    # Détection plus tolérante :
    # une ligne qui contient "poste" et une forme proche
    # de "intitulé".
    # --------------------------------------------------------

    if (
        "poste" in texte
        and (
            "linie" in texte
            or "linte" in texte
            or "intitul" in texte
            or "intitule" in texte
        )
    ):
        return True

    return False


# ============================================================
# LIBELLÉ CIBLE
# ============================================================

def _est_un_libelle_cible(ligne):

    return (
        _est_libelle_entreprise(ligne)
        or _est_libelle_poste(ligne)
        or _est_libelle_taches(ligne)
    )


# ============================================================
# NETTOYAGE VALEUR OCR
# ============================================================

def _nettoyer_valeur_ocr(valeur):

    if not valeur:
        return ""

    valeur = valeur.strip()

    valeur = re.sub(
        r"^[|¦]+",
        "",
        valeur
    )

    valeur = re.sub(
        r"[|¦]+$",
        "",
        valeur
    )

    return valeur.strip()


# ============================================================
# EXTRACTION APRÈS UN LIBELLÉ
# ============================================================

def _extraire_apres_texte(
    ligne,
    motif
):

    resultat = re.search(
        motif,
        ligne,
        flags=re.IGNORECASE
    )

    if not resultat:
        return ""

    valeur = ligne[resultat.end():]

    valeur = valeur.strip()

    valeur = valeur.lstrip(
        ":|¦ "
    )

    return _nettoyer_valeur_ocr(
        valeur
    )


# ============================================================
# EXTRACTION DU POSTE
# ============================================================

def _extraire_poste(
    lignes,
    index_poste
):

    # --------------------------------------------------------
    # 1. Valeur sur la même ligne
    # --------------------------------------------------------

    ligne = lignes[index_poste]

    motifs_poste = [
        r"intitul[ée]?\s+du\s+poste\s*:?",
        r"linie\s+du\s+poste\s*:?",
        r"linte\s+du\s+poste\s*:?",
        r"intitul\s+poste\s*:?",
    ]

    for motif in motifs_poste:

        valeur = _extraire_apres_texte(
            ligne,
            motif
        )

        if valeur:

            # Si la valeur contient une autre colonne,
            # on conserve uniquement la partie utile.
            valeur = valeur.split("|")[0].strip()

            if len(valeur) >= 3:

                return valeur

    # --------------------------------------------------------
    # 2. Valeur sur les lignes suivantes
    # --------------------------------------------------------

    for suivant in range(
        index_poste + 1,
        min(
            index_poste + 6,
            len(lignes)
        )
    ):

        candidat = lignes[suivant].strip()

        if not candidat:
            continue

        if _est_libelle_taches(candidat):
            continue

        if _est_libelle_entreprise(candidat):
            continue

        # Une autre rubrique clairement identifiée
        # peut arrêter la recherche.
        if (
            "habilitations" in candidat.lower()
            or "risques" in candidat.lower()
            or "conditions particulières" in candidat.lower()
        ):
            break

        morceaux = [
            _nettoyer_valeur_ocr(m)
            for m in candidat.split("|")
            if m.strip()
        ]

        # ----------------------------------------------------
        # Cas normal
        # ----------------------------------------------------

        if len(morceaux) == 1:

            valeur = morceaux[0]

            if len(valeur) >= 4:

                return valeur

        # ----------------------------------------------------
        # Cas colonne
        # ----------------------------------------------------

        for morceau in morceaux:

            morceau = morceau.strip()

            if len(morceau) >= 8:

                # On évite les textes manifestement
                # étrangers au poste.
                texte_bas = morceau.lower()

                if (
                    "habilitations" not in texte_bas
                    and "risques" not in texte_bas
                    and "conditions particulières" not in texte_bas
                    and "travaux sous" not in texte_bas
                ):

                    return morceau

    return ""


# ============================================================
# EXTRACTION DES TÂCHES
# ============================================================

def _extraire_taches(
    lignes,
    index_entete
):

    taches = []

    # --------------------------------------------------------
    # On inspecte seulement une petite zone après
    # l'en-tête du formulaire.
    #
    # Cela évite d'avaler tout le reste de la fiche.
    # --------------------------------------------------------

    limite = min(
        index_entete + 7,
        len(lignes)
    )

    for index in range(
        index_entete + 1,
        limite
    ):

        ligne = lignes[index].strip()

        if not ligne:
            continue

        # ----------------------------------------------------
        # Le poste indique la fin de la zone principale.
        # ----------------------------------------------------

        if _est_libelle_poste(ligne):

            break

        # ----------------------------------------------------
        # Découpage des colonnes OCR
        # ----------------------------------------------------

        morceaux = [
            _nettoyer_valeur_ocr(m)
            for m in ligne.split("|")
            if m.strip()
        ]

        # ----------------------------------------------------
        # On récupère les éléments qui ressemblent à des
        # tâches et pas aux libellés.
        # ----------------------------------------------------

        for morceau in morceaux:

            morceau = morceau.strip()

            if not morceau:
                continue

            texte_bas = morceau.lower()

            if (
                "nom de l'entreprise" in texte_bas
                or "liste des tâches proposées" in texte_bas
                or "intitulé du poste" in texte_bas
                or "intitule du poste" in texte_bas
            ):
                continue

            # ------------------------------------------------
            # Éléments manifestement parasites de l'OCR
            # ------------------------------------------------

            if (
                morceau.startswith("Dre panne")
                or morceau.startswith("!")
                or morceau.startswith("Sondtons")
                or morceau.startswith("Habilitalions")
                or morceau.startswith("FORCES")
            ):
                continue

            # ------------------------------------------------
            # Les vraies tâches de notre formulaire
            # sont des descriptions relativement courtes.
            # ------------------------------------------------

            if len(morceau) >= 4 and len(morceau) <= 100:

                taches.append(morceau)

    # --------------------------------------------------------
    # Suppression des doublons
    # --------------------------------------------------------

    taches_finales = []

    for tache in taches:

        if tache not in taches_finales:

            taches_finales.append(tache)

    return taches_finales


# ============================================================
# EXTRACTION CIBLÉE DE LA FICHE DE POSTE
# ============================================================

def extraire_fiche_poste_ciblee(texte):
    """
    Extrait :

    - entreprise
    - poste
    - tâches

    La fonction est adaptée aux formulaires OCR
    dont les rubriques sont organisées en colonnes.
    """

    resultat = {
        "entreprise": "",
        "poste": "",
        "taches": "",
        "entreprise_trouvee": False,
        "poste_trouve": False,
        "taches_trouvees": False,
    }

    if not texte:
        return resultat

    # --------------------------------------------------------
    # Préparation des lignes
    # --------------------------------------------------------

    lignes = texte.split("\n")

    lignes = [
        ligne.strip()
        for ligne in lignes
        if ligne.strip()
    ]

    # ========================================================
    # RECHERCHE DE L'EN-TÊTE DU FORMULAIRE
    # ========================================================

    index_entete = None

    for index, ligne in enumerate(lignes):

        normalisee = _normaliser_ligne(
            ligne
        )

        if (
            "nom de l'entreprise" in normalisee
            and "liste des tâches proposées" in normalisee
        ):

            index_entete = index
            break

    # ========================================================
    # ENTREPRISE
    # ========================================================

    if index_entete is not None:

        # ----------------------------------------------------
        # Recherche dans les lignes qui suivent l'en-tête
        # ----------------------------------------------------

        for index in range(
            index_entete + 1,
            min(
                index_entete + 4,
                len(lignes)
            )
        ):

            ligne = lignes[index]

            morceaux = [
                _nettoyer_valeur_ocr(m)
                for m in ligne.split("|")
                if m.strip()
            ]

            if not morceaux:
                continue

            # Le premier morceau correspond à la colonne
            # entreprise.
            premier = morceaux[0]

            texte_bas = premier.lower()

            if (
                premier
                and "nom de l'entreprise" not in texte_bas
                and "liste des tâches" not in texte_bas
                and len(premier) >= 2
            ):

                # On évite de prendre des morceaux manifestement
                # issus d'une autre rubrique.
                if (
                    not premier.startswith("Dre panne")
                    and not premier.startswith("!")
                ):

                    resultat["entreprise"] = premier
                    resultat["entreprise_trouvee"] = True
                    break

    # ========================================================
    # TÂCHES
    # ========================================================

    if index_entete is not None:

        taches = _extraire_taches(
            lignes,
            index_entete
        )

        # ----------------------------------------------------
        # Nettoyage supplémentaire :
        # on ne conserve que les tâches situées avant
        # l'intitulé du poste.
        # ----------------------------------------------------

        taches_nettoyees = []

        for tache in taches:

            texte_bas = tache.lower()

            if (
                "ouvrier vrd conducteur" in texte_bas
                or "poste" in texte_bas
                or "habilitation" in texte_bas
                or "risque" in texte_bas
                or "conditions particulières" in texte_bas
            ):
                continue

            if tache not in taches_nettoyees:

                taches_nettoyees.append(tache)

        # ----------------------------------------------------
        # Sur notre formulaire, les tâches utiles sont
        # les quatre premières descriptions.
        # ----------------------------------------------------

        if taches_nettoyees:

            taches_nettoyees = taches_nettoyees[:4]

            resultat["taches"] = ", ".join(
                taches_nettoyees
            )

            resultat["taches_trouvees"] = True

    # ========================================================
    # POSTE
    # ========================================================

    index_poste = None

    for index, ligne in enumerate(lignes):

        if _est_libelle_poste(ligne):

            index_poste = index
            break

    if index_poste is not None:

        poste = _extraire_poste(
            lignes,
            index_poste
        )

        if poste:

            resultat["poste"] = poste
            resultat["poste_trouve"] = True

    # ========================================================
    # SECOURS POUR LE POSTE
    # ========================================================

    # Si le libellé OCR est tellement déformé qu'il n'a
    # pas été reconnu, on recherche une ligne contenant
    # une forme proche de "poste".

    if not resultat["poste_trouve"]:

        for index, ligne in enumerate(lignes):

            texte_bas = _normaliser_ligne(
                ligne
            )

            if (
                "poste" in texte_bas
                and (
                    "linie" in texte_bas
                    or "linte" in texte_bas
                    or "intitul" in texte_bas
                )
            ):

                poste = _extraire_poste(
                    lignes,
                    index
                )

                if poste:

                    resultat["poste"] = poste
                    resultat["poste_trouve"] = True
                    break

    return resultat


# ============================================================
# GÉNÉRATION DE PRÉSENTATION CANDIDAT
# ============================================================

def generer_presentation(
    candidat,
    metier,
    competences,
    caces,
    permis,
    entreprise,
    agence
):
    """
    Génère une présentation d'un candidat destinée
    à l'entreprise cliente.
    """

    texte = f"""
Objet : Proposition de candidature – {metier}

Bonjour,

Suite à votre recherche, nous avons le plaisir de vous proposer la candidature de {candidat}.

Son profil présente plusieurs atouts :

- Métier : {metier}

- Compétences : {competences}

- CACES : {caces if caces else "Non renseigné"}

- Permis : {permis if permis else "Non renseigné"}

Ce candidat semble correspondre aux critères recherchés pour votre besoin.

Nous restons à votre disposition pour toute information complémentaire ou pour organiser une rencontre.

Cordialement,

ID'EES Intérim
Agence de {agence}
"""

    return texte.strip()
