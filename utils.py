import pdfplumber
import docx
import re
import pytesseract

from PIL import Image


# ============================================================
# EXTRACTION DE TEXTE
# PDF + WORD
# ============================================================

def extract_text(file):
    """
    Extrait le texte d'un fichier PDF ou Word.

    PDF :
    1. tentative de lecture classique avec pdfplumber ;
    2. extraction ciblée ;
    3. si la lecture classique est insuffisante ou mal structurée,
       OCR automatique avec Tesseract ;
    4. on conserve la meilleure version disponible.

    DOCX :
    - paragraphes ;
    - tableaux.
    """

    nom_fichier = getattr(file, "name", "") or ""
    nom_fichier_min = nom_fichier.lower()

    if nom_fichier_min.endswith(".docx"):

        texte = extraire_texte_docx(file)

        return nettoyer_texte(texte)

    elif nom_fichier_min.endswith(".pdf"):

        # ----------------------------------------------------
        # 1. Lecture classique
        # ----------------------------------------------------

        texte_classique = extraire_texte_pdf(file)

        texte_classique = nettoyer_texte(
            texte_classique
        )

        # ----------------------------------------------------
        # 2. Vérification de la lecture ciblée
        # ----------------------------------------------------

        fiche_classique = extraire_fiche_poste_ciblee(
            texte_classique
        )

        informations_classiques = sum(
            [
                bool(
                    fiche_classique.get(
                        "entreprise_trouvee"
                    )
                ),
                bool(
                    fiche_classique.get(
                        "poste_trouve"
                    )
                ),
                bool(
                    fiche_classique.get(
                        "taches_trouvees"
                    )
                ),
            ]
        )

        # ----------------------------------------------------
        # 3. OCR
        #
        # On lance volontairement l'OCR lorsque la lecture
        # classique ne permet pas de retrouver correctement
        # les rubriques importantes.
        #
        # Cela règle notamment les PDF modifiables dont
        # pdfplumber mélange les colonnes.
        # ----------------------------------------------------

        texte_ocr = ""

        if informations_classiques < 3:

            texte_ocr = extraire_texte_pdf_ocr(
                file
            )

            texte_ocr = nettoyer_texte(
                texte_ocr
            )

        # ----------------------------------------------------
        # 4. Comparaison des deux lectures
        # ----------------------------------------------------

        if texte_ocr:

            fiche_ocr = extraire_fiche_poste_ciblee(
                texte_ocr
            )

            informations_ocr = sum(
                [
                    bool(
                        fiche_ocr.get(
                            "entreprise_trouvee"
                        )
                    ),
                    bool(
                        fiche_ocr.get(
                            "poste_trouve"
                        )
                    ),
                    bool(
                        fiche_ocr.get(
                            "taches_trouvees"
                        )
                    ),
                ]
            )

            # ------------------------------------------------
            # L'OCR gagne si elle retrouve davantage
            # d'informations ciblées.
            # ------------------------------------------------

            if informations_ocr > informations_classiques:

                texte = texte_ocr

            else:

                texte = texte_classique

        else:

            texte = texte_classique

        return nettoyer_texte(
            texte
        )

    return ""


# ============================================================
# EXTRACTION PDF CLASSIQUE
# ============================================================

def extraire_texte_pdf(file):
    """
    Extraction classique avec pdfplumber.

    On utilise plusieurs stratégies car certains PDF
    modifiables sont construits avec plusieurs colonnes.
    """

    morceaux = []

    try:

        try:
            file.seek(0)
        except Exception:
            pass

        with pdfplumber.open(file) as pdf:

            for page in pdf.pages:

                texte_page = ""

                # ------------------------------------------------
                # Première tentative
                # ------------------------------------------------

                try:

                    texte_page = page.extract_text(
                        x_tolerance=2,
                        y_tolerance=3,
                    ) or ""

                except Exception:

                    texte_page = ""

                # ------------------------------------------------
                # Deuxième tentative si nécessaire
                # ------------------------------------------------

                if not texte_page.strip():

                    try:

                        texte_page = page.extract_text(
                            x_tolerance=1,
                            y_tolerance=2,
                        ) or ""

                    except Exception:

                        texte_page = ""

                if texte_page:

                    morceaux.append(
                        texte_page
                    )

    except Exception:

        return ""

    return "\n".join(
        morceaux
    )


# ============================================================
# EXTRACTION PDF PAR OCR
# ============================================================

def extraire_texte_pdf_ocr(file):
    """
    Convertit chaque page PDF en image puis utilise Tesseract.

    Le français est privilégié.
    """

    morceaux = []

    try:

        try:
            file.seek(0)
        except Exception:
            pass

        with pdfplumber.open(file) as pdf:

            for page in pdf.pages:

                try:

                    image_page = page.to_image(
                        resolution=300
                    ).original

                    texte_page = pytesseract.image_to_string(
                        image_page,
                        lang="fra+eng",
                        config="--psm 6",
                    )

                    if texte_page:

                        morceaux.append(
                            texte_page
                        )

                except Exception:

                    continue

    except Exception:

        return ""

    return "\n".join(
        morceaux
    )


# ============================================================
# EXTRACTION WORD
# ============================================================

def extraire_texte_docx(file):
    """
    Extrait le texte d'un document Word.

    Les paragraphes et les tableaux sont lus.
    """

    morceaux = []

    try:

        try:
            file.seek(0)
        except Exception:
            pass

        document = docx.Document(
            file
        )

        # ----------------------------------------------------
        # Paragraphes
        # ----------------------------------------------------

        for paragraphe in document.paragraphs:

            texte = paragraphe.text.strip()

            if texte:

                morceaux.append(
                    texte
                )

        # ----------------------------------------------------
        # Tableaux
        # ----------------------------------------------------

        for table in document.tables:

            for ligne in table.rows:

                cellules = []

                for cellule in ligne.cells:

                    texte_cellule = (
                        cellule.text.strip()
                    )

                    if texte_cellule:

                        cellules.append(
                            texte_cellule
                        )

                if cellules:

                    morceaux.append(
                        " | ".join(
                            cellules
                        )
                    )

    except Exception:

        return ""

    return "\n".join(
        morceaux
    )


# ============================================================
# NETTOYAGE
# ============================================================

def nettoyer_texte(texte):
    """
    Nettoyage général.

    IMPORTANT :
    Les retours à la ligne sont conservés.
    """

    if not texte:

        return ""

    texte = texte.replace(
        "\r\n",
        "\n"
    )

    texte = texte.replace(
        "\r",
        "\n"
    )

    texte = texte.replace(
        "\xa0",
        " "
    )

    # --------------------------------------------------------
    # Supprimer les liens SVG accidentellement récupérés
    # --------------------------------------------------------

    texte = re.sub(
        r"\[svg\]\([^)]+\)",
        "",
        texte,
        flags=re.IGNORECASE,
    )

    # --------------------------------------------------------
    # Espaces multiples
    # --------------------------------------------------------

    texte = re.sub(
        r"[ \t]+",
        " ",
        texte,
    )

    # --------------------------------------------------------
    # Lignes vides multiples
    # --------------------------------------------------------

    texte = re.sub(
        r"\n[ \t]*\n[ \t]*\n+",
        "\n\n",
        texte,
    )

    return texte.strip()


# ============================================================
# NORMALISATION
# ============================================================

def _normaliser_texte_recherche(texte):
    """
    Normalisation destinée uniquement à la recherche
    des rubriques.

    Les accents sont supprimés afin de mieux résister
    aux différences entre PDF et OCR.
    """

    if not texte:

        return ""

    texte = texte.lower()

    texte = texte.replace(
        "’",
        "'"
    )

    texte = texte.replace(
        "|",
        " "
    )

    texte = texte.replace(
        ":", 
        " "
    )

    # Accents français fréquents
    remplacements = {
        "à": "a",
        "â": "a",
        "ä": "a",
        "é": "e",
        "è": "e",
        "ê": "e",
        "ë": "e",
        "î": "i",
        "ï": "i",
        "ô": "o",
        "ö": "o",
        "ù": "u",
        "û": "u",
        "ü": "u",
        "ç": "c",
    }

    for ancien, nouveau in remplacements.items():

        texte = texte.replace(
            ancien,
            nouveau
        )

    texte = re.sub(
        r"\s+",
        " ",
        texte
    )

    return texte.strip()


def _normaliser_ligne(ligne):

    if not ligne:

        return ""

    ligne = ligne.strip()

    ligne = ligne.replace(
        "’",
        "'"
    )

    ligne = ligne.replace(
        "|",
        " "
    )

    ligne = ligne.replace(
        "®",
        " "
    )

    ligne = re.sub(
        r"\s+",
        " ",
        ligne
    )

    return ligne.strip()


# ============================================================
# DETECTION DES RUBRIQUES
# ============================================================

def _position_libelle_entreprise(texte):

    if not texte:

        return None

    texte_recherche = _normaliser_texte_recherche(
        texte
    )

    motifs = [
        r"nom\s+de\s+l[' ]?entreprise",
        r"nom\s+entreprise",
    ]

    for motif in motifs:

        resultat = re.search(
            motif,
            texte_recherche,
            flags=re.IGNORECASE,
        )

        if resultat:

            return (
                resultat.start(),
                resultat.end()
            )

    return None


def _position_libelle_poste(texte):

    if not texte:

        return None

    texte_recherche = _normaliser_texte_recherche(
        texte
    )

    motifs = [
        r"intitule\s+du\s+poste",
        r"intitul\s+du\s+poste",
        r"intitule\s+de\s+poste",
        r"intitul\s+de\s+poste",
        r"intitule\s+poste",
        r"intitul\s+poste",
    ]

    for motif in motifs:

        resultat = re.search(
            motif,
            texte_recherche,
            flags=re.IGNORECASE,
        )

        if resultat:

            return (
                resultat.start(),
                resultat.end()
            )

    return None


def _position_libelle_taches(texte):

    if not texte:

        return None

    texte_recherche = _normaliser_texte_recherche(
        texte
    )

    motifs = [
        r"liste\s+des\s+taches\s+proposees",
        r"liste\s+des\s+taches\s+proposee",
        r"liste\s+taches\s+proposees",
        r"liste\s+taches",
    ]

    for motif in motifs:

        resultat = re.search(
            motif,
            texte_recherche,
            flags=re.IGNORECASE,
        )

        if resultat:

            return (
                resultat.start(),
                resultat.end()
            )

    return None


# ============================================================
# TEST DES RUBRIQUES
# ============================================================

def _est_libelle_entreprise(ligne):

    return (
        _position_libelle_entreprise(
            ligne
        )
        is not None
    )


def _est_libelle_poste(ligne):

    return (
        _position_libelle_poste(
            ligne
        )
        is not None
    )


def _est_libelle_taches(ligne):

    return (
        _position_libelle_taches(
            ligne
        )
        is not None
    )


def _est_un_libelle_cible(ligne):

    return (
        _est_libelle_entreprise(ligne)
        or _est_libelle_poste(ligne)
        or _est_libelle_taches(ligne)
    )


# ============================================================
# NETTOYAGE D'UNE VALEUR
# ============================================================

def _nettoyer_valeur(valeur):

    if not valeur:

        return ""

    valeur = valeur.strip()

    valeur = valeur.strip(
        "|:;,-"
    )

    valeur = re.sub(
        r"\s+",
        " ",
        valeur
    )

    return valeur.strip()


# ============================================================
# VALEUR DE RUBRIQUE SUR UNE LIGNE
# ============================================================

def _valeur_sur_meme_ligne(
    ligne,
    type_information
):

    if not ligne:

        return ""

    if type_information == "entreprise":

        motifs = [
            r"nom\s+de\s+l[' ]?entreprise\s*:?\s*(.+)$",
            r"nom\s+entreprise\s*:?\s*(.+)$",
        ]

    elif type_information == "poste":

        motifs = [
            r"intitul[ée]?\s+du\s+poste\s*:?\s*(.+)$",
            r"intitul[ée]?\s+de\s+poste\s*:?\s*(.+)$",
            r"intitul[ée]?\s+poste\s*:?\s*(.+)$",
        ]

    else:

        motifs = [
            r"liste\s+des\s+t[âa]ches\s+propos[ée]es\s*:?\s*(.+)$",
            r"liste\s+des\s+taches\s+proposees\s*:?\s*(.+)$",
        ]

    for motif in motifs:

        resultat = re.search(
            motif,
            ligne,
            flags=re.IGNORECASE,
        )

        if resultat:

            valeur = resultat.group(
                1
            )

            return _nettoyer_valeur(
                valeur
            )

    return ""


# ============================================================
# EXTRACTION GENERIQUE D'UNE RUBRIQUE
# ============================================================

def _extraire_rubrique(
    texte,
    type_information
):

    if not texte:

        return ""

    lignes = texte.split(
        "\n"
    )

    # --------------------------------------------------------
    # 1. Recherche ligne par ligne
    # --------------------------------------------------------

    for index, ligne_originale in enumerate(
        lignes
    ):

        ligne = _normaliser_ligne(
            ligne_originale
        )

        if not ligne:

            continue

        # ----------------------------------------------------
        # Détection du libellé
        # ----------------------------------------------------

        if type_information == "entreprise":

            est_libelle = _est_libelle_entreprise(
                ligne
            )

        elif type_information == "poste":

            est_libelle = _est_libelle_poste(
                ligne
            )

        else:

            est_libelle = _est_libelle_taches(
                ligne
            )

        if not est_libelle:

            continue

        # ----------------------------------------------------
        # 2. Valeur sur la même ligne
        # ----------------------------------------------------

        valeur = _valeur_sur_meme_ligne(
            ligne,
            type_information
        )

        # ----------------------------------------------------
        # Si la ligne contient plusieurs rubriques,
        # on vérifie que la valeur ne contient pas
        # le libellé d'une autre rubrique.
        # ----------------------------------------------------

        if valeur:

            valeur_min = _normaliser_texte_recherche(
                valeur
            )

            if (
                "liste des taches" in valeur_min
                or "intitule du poste" in valeur_min
                or "nom de l entreprise" in valeur_min
            ):

                valeur = ""

            else:

                # Pour entreprise/poste :
                # une valeur courte et propre suffit.

                if type_information in (
                    "entreprise",
                    "poste",
                ):

                    return valeur

        # ----------------------------------------------------
        # 3. Recherche dans les lignes suivantes
        # ----------------------------------------------------

        valeurs = []

        for suivante in lignes[
            index + 1:
        ]:

            suivante = _normaliser_ligne(
                suivante
            )

            if not suivante:

                continue

            # Une autre rubrique cible arrête la recherche.

            if _est_un_libelle_cible(
                suivante
            ):

                break

            # ------------------------------------------------
            # ENTREPRISE
            # ------------------------------------------------

            if type_information == "entreprise":

                valeur_suivante = _nettoyer_valeur(
                    suivante
                )

                if valeur_suivante:

                    return valeur_suivante

            # ------------------------------------------------
            # POSTE
            # ------------------------------------------------

            elif type_information == "poste":

                valeur_suivante = _nettoyer_valeur(
                    suivante
                )

                if valeur_suivante:

                    return valeur_suivante

            # ------------------------------------------------
            # TACHES
            # ------------------------------------------------

            elif type_information == "taches":

                valeurs.append(
                    _nettoyer_valeur(
                        suivante
                    )
                )

        if type_information == "taches":

            valeurs_propres = []

            for valeur_tache in valeurs:

                if not valeur_tache:

                    continue

                # Éléments manifestement parasites
                if valeur_tache in (
                    "=",
                    "!",
                    "-",
                    "_",
                ):

                    continue

                valeurs_propres.append(
                    valeur_tache
                )

            if valeurs_propres:

                return ", ".join(
                    valeurs_propres
                )

    # --------------------------------------------------------
    # 4. Recherche de secours dans le texte complet
    # --------------------------------------------------------

    if type_information == "entreprise":

        position = _position_libelle_entreprise(
            texte
        )

    elif type_information == "poste":

        position = _position_libelle_poste(
            texte
        )

    else:

        position = _position_libelle_taches(
            texte
        )

    if position is None:

        return ""

    reste = texte[
        position[1]:
    ]

    # --------------------------------------------------------
    # On coupe au prochain libellé cible.
    # --------------------------------------------------------

    positions = []

    for fonction in (
        _position_libelle_entreprise,
        _position_libelle_poste,
        _position_libelle_taches,
    ):

        pos = fonction(
            reste
        )

        if pos:

            positions.append(
                pos[0]
            )

    if positions:

        reste = reste[
            :min(positions)
        ]

    lignes_reste = reste.split(
        "\n"
    )

    valeurs = []

    for ligne in lignes_reste:

        ligne = _normaliser_ligne(
            ligne
        )

        ligne = _nettoyer_valeur(
            ligne
        )

        if not ligne:

            continue

        if _est_un_libelle_cible(
            ligne
        ):

            break

        if ligne in (
            "=",
            "!",
            "-",
            "_",
        ):

            continue

        valeurs.append(
            ligne
        )

        if type_information in (
            "entreprise",
            "poste",
        ):

            break

    if type_information == "taches":

        return ", ".join(
            valeurs
        )

    if valeurs:

        return valeurs[0]

    return ""


# ============================================================
# EXTRACTION ENTREPRISE
# ============================================================

def _extraire_entreprise_depuis_texte(
    texte
):

    return _extraire_rubrique(
        texte,
        "entreprise"
    )


# ============================================================
# EXTRACTION POSTE
# ============================================================

def _extraire_poste_depuis_texte(
    texte
):

    return _extraire_rubrique(
        texte,
        "poste"
    )


# ============================================================
# EXTRACTION TACHES
# ============================================================

def _extraire_taches_depuis_texte(
    texte
):

    return _extraire_rubrique(
        texte,
        "taches"
    )


# ============================================================
# EXTRACTION CIBLEE
# ============================================================

def extraire_fiche_poste_ciblee(
    texte
):
    """
    Extrait uniquement :

    - entreprise
    - poste
    - tâches

    Ne déduit aucune information absente.
    """

    resultat = {
        "entreprise": "",
        "poste": "",
        "taches": "",
        "entreprise_trouvee": False,
        "poste_trouve": False,
        "taches_trouvees": False,
    }

    if not texte:

        return resultat

    texte_normalise = nettoyer_texte(
        texte
    )

    # --------------------------------------------------------
    # ENTREPRISE
    # --------------------------------------------------------

    entreprise = _extraire_entreprise_depuis_texte(
        texte_normalise
    )

    if entreprise:

        resultat["entreprise"] = entreprise
        resultat["entreprise_trouvee"] = True

    # --------------------------------------------------------
    # POSTE
    # --------------------------------------------------------

    poste = _extraire_poste_depuis_texte(
        texte_normalise
    )

    if poste:

        resultat["poste"] = poste
        resultat["poste_trouve"] = True

    # --------------------------------------------------------
    # TACHES
    # --------------------------------------------------------

    taches = _extraire_taches_depuis_texte(
        texte_normalise
    )

    if taches:

        resultat["taches"] = taches
        resultat["taches_trouvees"] = True

    # --------------------------------------------------------
    # SECURITE
    # --------------------------------------------------------

    valeurs_interdites = {
        "nom de l'entreprise",
        "nom de l entreprise",
        "nom entreprise",
        "liste des tâches proposées",
        "liste des taches proposées",
        "liste des taches proposees",
        "liste des taches",
        "intitulé du poste",
        "intitule du poste",
        "intitul du poste",
        "intitule poste",
        "intitul poste",
    }

    for cle in (
        "entreprise",
        "poste",
        "taches",
    ):

        valeur = (
            resultat[cle]
            or ""
        )

        valeur_min = (
            _normaliser_texte_recherche(
                valeur
            )
        )

        if valeur_min in valeurs_interdites:

            resultat[cle] = ""

            resultat[
                f"{cle}_trouvee"
            ] = False

    return resultat


# ============================================================
# GENERATION PRESENTATION CANDIDAT
# ============================================================

def generer_presentation(
    candidat,
    metier,
    competences,
    caces,
    permis,
    entreprise,
    agence
):
    """
    Génère une présentation d'un candidat
    destinée à l'entreprise cliente.
    """

    texte = f"""
Objet : Proposition de candidature – {metier}

Bonjour,

Suite à votre recherche, nous avons le plaisir de vous proposer la candidature de {candidat}.

Son profil présente plusieurs atouts :

- Métier : {metier}

- Compétences : {competences}

- CACES : {caces if caces else "Non renseigné"}

- Permis : {permis if permis else "Non renseigné"}

Ce candidat semble correspondre aux critères recherchés pour votre besoin.

Nous restons à votre disposition pour toute information complémentaire ou pour organiser une rencontre.

Cordialement,

ID'EES Intérim
Agence de {agence}
"""

    return texte.strip()
